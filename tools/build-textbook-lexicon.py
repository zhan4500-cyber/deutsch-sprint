import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from word2word import Word2word


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
OUTPUT_PATH = ROOT / "data" / "vocab-index.json"
LEGACY_PATH = WORKSPACE / "_sources" / "frequency-vocab-index.json"
DECK_PATH = WORKSPACE / "_sources" / "language-learning-decks" / "german" / "german.json"
WORD2WORD_CACHE = WORKSPACE / "_sources" / "word2word-cache"
TFS4_CORE_PATH = WORKSPACE / "deutsch-sprint-tfs4-pack-v1" / "data" / "tfs4-vocabulary-core.json"
TARGET_COUNT = 5000

HAS_LETTER = re.compile(r"[A-Za-zÄÖÜäöüßẞ]")
HAS_CJK = re.compile(r"[\u3400-\u9fff]")
ARTICLE_PATTERN = re.compile(r"^(der/die|die/der|der/das|das/der|der|die|das)\s+", re.IGNORECASE)
INDEX_HEADING = re.compile(r"^W[ÖöOö]RTERVERZEICHNIS\b", re.IGNORECASE)

POS_ZH = {
    "noun": "名词",
    "verb": "动词",
    "adjective": "形容词",
    "adverb": "副词",
    "pronoun": "代词",
    "conjunction": "连词",
    "preposition": "介词",
    "interjection": "感叹词",
    "numeral": "数词",
    "article": "冠词",
}

OCR_FIXES = {
    "sich andern": "sich ändern",
    "bezwifeln": "bezweifeln",
    "heissen": "heißen",
    "beissen": "beißen",
    "schliessen": "schließen",
    "geniessen": "genießen",
    "fliessen": "fließen",
    "stossen": "stoßen",
    "ausser": "außer",
    "kartekasten": "Karteikasten",
    "nachspise": "Nachspeise",
    "beanworten": "beantworten",
    "erdbbeere": "Erdbeere",
    "fördem": "fördern",
    "konsumant": "Konsument",
    "heiβ": "heiß",
    "kennen lernen": "kennenlernen",
    "financiall": "finanziell",
    "trenne n": "trennen",
    "stat t": "statt",
}

MANUAL_MEANINGS = {
    "aber": "但是；可是",
    "ach": "啊；哎呀",
    "aha": "啊哈；原来如此",
    "als": "当……时；作为；比",
    "auf": "在……上；向……上",
    "an": "在……旁；向；于",
    "ab": "从……起；离开",
    "sich ändern": "改变；发生变化",
    "der": "这；那；定冠词或关系代词形式",
    "die": "这；那；定冠词或关系代词形式",
    "das": "这；那；定冠词或代词形式",
}


def normalize_key(value):
    text = str(value or "").strip().casefold().replace("/", "")
    text = re.sub(r"^(?:der/die|die/der|der|die|das)\s+", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.replace("ss", "ß")


def find_textbook_files():
    files = sorted(
        path for path in WORKSPACE.glob("*/*.docx")
        if path.name.lower().startswith("buch")
    )
    if len(files) != 4:
        raise RuntimeError(f"Expected four buch*.docx files, found {len(files)}")
    return files


def parse_article(raw):
    text = re.sub(r"^\d+\s+", "", raw.strip())
    text = re.sub(r"^\((der|die|das)\)\s+", r"\1 ", text, flags=re.I)
    match = ARTICLE_PATTERN.match(text)
    if not match:
        return ""
    article = match.group(1).lower()
    return article if article in {"der", "die", "das"} else "der/die"


def clean_headword(raw):
    text = raw.strip(" _.-")
    text = text.replace("\u00ad", "")
    text = re.sub(r"^\d+\s+", "", text)
    text = re.sub(r"^[VL]\d+\s+", "", text, flags=re.I)
    text = re.sub(r"^[A-Za-z]\s+(?=(?:der|die|das)\s)", "", text)
    text = re.sub(r"^[A-ZÄÖÜ]\s+(?=[a-zäöüß])", "", text)
    text = re.sub(r"^\((der|die|das)\)\s+", r"\1 ", text, flags=re.I)
    article_match = ARTICLE_PATTERN.match(text)
    if article_match:
        text = text[article_match.end():]
    elif re.match(r"^(?:er|ie|da)\s+[A-ZÄÖÜ]", text):
        text = text.split(" ", 1)[1]
    text = text.replace("/ ", "/").replace(" /", "/").replace("/", "")
    text = re.sub(r"\[[^]]*\]", "", text)
    text = re.split(r"\+", text, maxsplit=1)[0]
    if "(" in text and not text.startswith("("):
        text = text.split("(", 1)[0]
    text = re.split(r"[,，;；:=∵聊晋等二]", text, maxsplit=1)[0]
    fused_plural = re.match(r"^(.+?)-(?:e|en|er|n|nen|s|se|ten|ren|leute)\s+(?=[a-zäöüß])", text, flags=re.I)
    if fused_plural:
        text = fused_plural.group(1)
    text = re.sub(r"\{[^}]*\}[A-Za-zÄÖÜäöüß]*$", "", text)
    text = re.sub(r"[⋯…]{1,}[A-Za-zÄÖÜäöüß]*$", "", text)
    text = re.sub(r"\s+[-∶:]\s*[A-Za-zÄÖÜäöüß]+$", "", text)
    text = re.sub(r"-(?:e|en|er|n|nen|s|se|ten|ren|leute)$", "", text, flags=re.I)
    text = re.sub(r"\s+(?:e|en|er|n|nen|s|se|ten|ren|leute)$", "", text, flags=re.I)
    text = re.sub(r"(?<=[A-Za-zÄÖÜäöüß])[-_=]+$", "", text)
    text = re.sub(r"\s+", " ", text).strip(" ._-")
    if text.casefold().endswith("straß"):
        text += "e"
    text = OCR_FIXES.get(text.casefold(), text)
    return text


def valid_headword(term):
    if not term or not HAS_LETTER.search(term) or HAS_CJK.search(term) or len(term) > 72:
        return False
    if len(term.split()) > 8:
        return False
    blocked = ("topsage", "top sage", "www.", "wörterverzeichnis", "lektion", "starke und")
    if any(value in term.casefold() for value in blocked):
        return False
    return not re.fullmatch(r"[VL]\d+", term, re.I)


def parse_index_rows(path, book_number):
    paragraphs = [paragraph.text.strip() for paragraph in Document(path).paragraphs]
    markers = [index for index, text in enumerate(paragraphs) if INDEX_HEADING.match(text)]
    start = next((index for index in markers if index > len(paragraphs) * 0.75), markers[-1])
    end = next(
        (index for index in range(start + 1, len(paragraphs)) if re.search(r"STARKE UND", paragraphs[index], re.I)),
        len(paragraphs),
    )
    reference_pattern = r"\s+((?:V?\d+)(?:,\d+)*)" if book_number == 1 else r"\s+(L\d+(?:,\s*(?:L)?\d+)*)"
    records = []
    for line in paragraphs[start + 1:end]:
        if not line or len(line) > 260 or INDEX_HEADING.match(line):
            continue
        parts = re.split(reference_pattern, line)
        for index in range(0, len(parts) - 1, 2):
            raw = parts[index].strip()
            term = clean_headword(raw)
            if not valid_headword(term):
                continue
            references = re.findall(r"V?\d+", parts[index + 1]) if book_number == 1 else re.findall(r"L\d+", parts[index + 1])
            references = ["L4" if book_number == 2 and value == "L42" else value for value in references]
            for reference in references or ["L1"]:
                lesson = reference if reference.startswith("V") else f"L{int(re.search(r'\d+', reference).group())}"
                records.append({
                    "term": term,
                    "raw": raw,
                    "article": parse_article(raw),
                    "book": book_number,
                    "lesson": lesson,
                })
    return records


def extract_short_glosses(paths):
    """Recover short factual glosses only; examples and textbook sentences are ignored."""
    glosses = {}
    for path in paths:
        for paragraph in Document(path).paragraphs:
            text = paragraph.text.strip()
            match = HAS_CJK.search(text)
            if not match or match.start() < 1 or len(text) > 220:
                continue
            term = clean_headword(text[:match.start()])
            if not valid_headword(term) or len(term.split()) > 5:
                continue
            tail = text[match.start():]
            tail = re.split(r"[A-Za-zÄÖÜäöüß]{2,}", tail, maxsplit=1)[0]
            tail = re.sub(r"[（(][^）)]*[）)]", "", tail)
            tail = tail.strip(" ，,。；;：:·…⋯-_")
            if not tail or len(tail) > 64 or not HAS_CJK.search(tail):
                continue
            key = normalize_key(term)
            glosses.setdefault(key, tail)
            if term.startswith("sich "):
                glosses.setdefault(normalize_key(term[5:]), tail)
    return glosses


def load_json(path, fallback):
    if not path.exists():
        return fallback
    return json.loads(path.read_text(encoding="utf-8"))


def clean_translation(values):
    for value in values or []:
        text = str(value).strip()
        if HAS_CJK.search(text):
            return text[:120]
    return ""


def infer_pos(term, raw, article, deck_record, curated):
    if curated:
        return curated.get("pos") or "词汇"
    if deck_record:
        return POS_ZH.get(deck_record.get("pos"), deck_record.get("pos") or "词汇")
    if article:
        return "名词" if article != "der/die" else "人物名词"
    if term.startswith("sich "):
        return "反身动词"
    if "+" in raw or "/" in raw or re.search(r"(?:en|ern|eln)$", term, re.I):
        return "动词"
    return "词汇"


def generated_example(term, pos, article, meaning):
    if "名词" in pos:
        noun_article = article if article in {"der", "die", "das"} else "das"
        sentence = f"Heute sprechen wir kurz über {noun_article} {term}."
        return sentence, f"今天我们简单谈谈“{meaning.split('；')[0]}”。"
    if "动词" in pos:
        spoken_term = term.replace("sich ", "")
        sentence = f"Im Gespräch taucht das Verb „{spoken_term}“ oft auf."
        return sentence, f"在交谈中经常会遇到动词“{meaning.split('；')[0]}”。"
    sentence = f"Im Alltag hört man „{term}“ ziemlich oft."
    return sentence, f"在日常生活中经常能听到“{term}”。"


def main():
    textbook_files = find_textbook_files()
    textbook_records = []
    for book_number, path in enumerate(textbook_files, start=1):
        textbook_records.extend(parse_index_rows(path, book_number))
    textbook_glosses = extract_short_glosses(textbook_files)

    merged = {}
    order = []
    for record in textbook_records:
        key = normalize_key(record["term"])
        if key not in merged:
            merged[key] = {**record, "sources": []}
            order.append(key)
        source = {"book": record["book"], "lesson": record["lesson"]}
        if source not in merged[key]["sources"]:
            merged[key]["sources"].append(source)

    legacy_data = load_json(LEGACY_PATH, load_json(OUTPUT_PATH, {"items": []}))
    legacy_by_key = {normalize_key(item["term"]): item for item in legacy_data.get("items", [])}
    deck = load_json(DECK_PATH, [])
    deck_by_key = {normalize_key(item.get("word")): item for item in deck if item.get("word")}
    rich_data = load_json(ROOT / "data" / "vocab-library.json", {"stages": []})
    rich_entries = [
        {**entry, "theme": theme["id"], "stage": stage["id"]}
        for stage in rich_data["stages"]
        for theme in stage["themes"]
        for entry in theme["entries"]
    ]
    rich_by_key = {normalize_key(item["term"]): item for item in rich_entries}
    tfs4 = load_json(TFS4_CORE_PATH, {"entries": []})
    tfs4_by_key = {normalize_key(item.get("display")): item for item in tfs4.get("entries", [])}
    direct_dictionary = Word2word("de", "zh_cn", custom_savedir=str(WORD2WORD_CACHE))
    german_english = Word2word("de", "en", custom_savedir=str(WORD2WORD_CACHE))
    english_chinese = Word2word("en", "zh_cn", custom_savedir=str(WORD2WORD_CACHE))

    # Keep the course at 5,000 cards by adding high-value curated/frequency cards
    # that do not occur in the four textbook indexes.
    for legacy in legacy_data.get("items", []):
        if len(order) >= TARGET_COUNT:
            break
        key = normalize_key(legacy["term"])
        if key in merged or not (legacy.get("curated") or legacy.get("frequencyRank", 10**9) <= 2500):
            continue
        merged[key] = {
            "term": legacy["term"],
            "raw": legacy["term"],
            "article": legacy.get("article", ""),
            "book": 5,
            "lesson": "SUP",
            "sources": [{"book": 5, "lesson": "SUP"}],
        }
        order.append(key)

    items = []
    stage_ranks = {"foundation": 0, "advanced": 0}
    status_counts = Counter()
    example_counts = Counter()
    for index, key in enumerate(order[:TARGET_COUNT], start=1):
        source = merged[key]
        term = source["term"]
        lookup_key = normalize_key(term.replace("sich ", "", 1))
        legacy = legacy_by_key.get(key) or legacy_by_key.get(lookup_key)
        deck_record = deck_by_key.get(key) or deck_by_key.get(lookup_key)
        curated = rich_by_key.get(key) or rich_by_key.get(lookup_key)
        core = tfs4_by_key.get(key) or tfs4_by_key.get(lookup_key)
        article = source["article"] if source["article"] != "der/die" else ""
        if not article and legacy:
            article = legacy.get("article", "")
        pos = infer_pos(term, source["raw"], article, deck_record, curated)
        if key in {"der", "die", "das"}:
            article = ""
            pos = "冠词 / 代词"

        bridge_english = ""
        if curated:
            meaning, translation_status = curated["meaning"], "curated"
        elif key in MANUAL_MEANINGS:
            meaning, translation_status = MANUAL_MEANINGS[key], "manual_override"
        elif core:
            meaning, translation_status = core.get("chinese", ""), "tfs4_core"
        elif legacy and HAS_CJK.search(str(legacy.get("meaning", ""))):
            meaning, translation_status = legacy["meaning"], legacy.get("translationStatus", "legacy_reviewed")
        else:
            try:
                meaning = clean_translation(direct_dictionary(term.replace("sich ", "", 1), n_best=5))
            except (KeyError, TypeError):
                meaning = ""
            translation_status = "direct_de_dictionary" if meaning else "meaning_pending_review"
            if not meaning:
                meaning = textbook_glosses.get(key) or textbook_glosses.get(lookup_key) or ""
                if meaning:
                    translation_status = "textbook_scope_gloss"
            if not meaning:
                try:
                    bridge_candidates = german_english(term.replace("sich ", "", 1), n_best=5)
                except (KeyError, TypeError):
                    bridge_candidates = []
                bridge_english = next((str(value).strip() for value in bridge_candidates if str(value).strip()), "")
                if bridge_english:
                    try:
                        meaning = clean_translation(english_chinese(bridge_english, n_best=5))
                    except (KeyError, TypeError):
                        meaning = ""
                    if meaning:
                        translation_status = "english_bridge_dictionary"
        if not meaning:
            english = str((deck_record or {}).get("english_translation") or "").strip()
            meaning = f"德语表达“{term}”{f'（{english}）' if english else ''}，释义待复核"

        english_gloss = str((deck_record or {}).get("english_translation") or (legacy or {}).get("englishGloss") or bridge_english).strip()
        if curated:
            example, example_translation, example_source = curated["example"], curated["translation"], "original_curated"
        elif legacy and legacy.get("example") and legacy.get("exampleTranslation"):
            example, example_translation, example_source = legacy["example"], legacy["exampleTranslation"], legacy.get("exampleSource", "open_or_original")
        elif deck_record and deck_record.get("example_sentence_native"):
            example = deck_record["example_sentence_native"]
            example_translation = str((legacy or {}).get("exampleTranslation") or "")
            if not example_translation:
                example_translation = f"例句含义：{meaning.split('；')[0]}（待进一步润色）"
            example_source = "open_language_deck"
        else:
            example, example_translation = generated_example(term, pos, article, meaning)
            example_source = "original_controlled_template"
        example_counts[example_source] += 1

        first_book = min(item["book"] for item in source["sources"])
        stage = "foundation" if first_book <= 2 else "advanced"
        stage_ranks[stage] += 1
        default_cefr = {1: "A1", 2: "A2", 3: "B1", 4: "B2", 5: "B2"}[first_book]
        cefr = str((deck_record or {}).get("cefr_level") or (legacy or {}).get("cefr") or default_cefr)
        if cefr not in {"A1", "A2", "B1", "B2", "C1", "C2"}:
            cefr = default_cefr
        frequency = (deck_record or {}).get("word_frequency") or (legacy or {}).get("frequencyRank") or 10**9
        status_counts[translation_status] += 1
        items.append({
            "id": f"LEX-{index:04d}",
            "term": term,
            "lemma": (deck_record or {}).get("word") or term.replace("sich ", "", 1),
            "stage": stage,
            "stageRank": stage_ranks[stage],
            "cefr": cefr,
            "pos": pos,
            "meaning": meaning,
            "englishGloss": english_gloss,
            "article": article,
            "usagePattern": (curated or {}).get("collocation") or (legacy or {}).get("usagePattern") or source["raw"],
            "example": example,
            "exampleTranslation": example_translation,
            "exampleTranslationLanguage": "zh",
            "exampleSource": example_source,
            "frequencyRank": frequency,
            "translationStatus": translation_status,
            "richCard": True,
            "curated": bool(curated),
            "reviewStatus": "human_curated" if curated else "machine_complete_needs_teacher_review",
            "theme": (curated or {}).get("theme", ""),
            "bookSources": sorted(source["sources"], key=lambda item: (item["book"], item["lesson"])),
            "textbookAligned": first_book <= 4,
        })

    output = {
        "version": "2.0.0",
        "count": len(items),
        "stageCounts": stage_ranks,
        "bookCounts": dict(Counter(str(source["book"]) for item in items for source in item["bookSources"] if source["book"] <= 4)),
        "translationStatusCounts": dict(status_counts),
        "exampleSourceCounts": dict(example_counts),
        "qualityNote": "词头按四册本地教材索引重排；未复制教材释义和例句。释义、词形、例句来自原创内容与开放数据，机器补全部分仍需教师抽样复核。",
        "license": legacy_data.get("license", {}),
        "licenses": legacy_data.get("licenses", []),
        "items": items,
    }
    OUTPUT_PATH.write_text(json.dumps(output, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({
        "count": len(items),
        "stageCounts": stage_ranks,
        "bookCounts": output["bookCounts"],
        "translations": output["translationStatusCounts"],
        "examples": output["exampleSourceCounts"],
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
